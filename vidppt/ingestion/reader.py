"""将 Word/PDF 教案读取为保留标题层级的 SourceDocument。"""

from __future__ import annotations

from pathlib import Path

from .models import SourceDocument, SourceSection


def read_source_document(path: Path) -> SourceDocument:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    raise ValueError(f"教案输入仅支持 .docx 和 .pdf，收到: {suffix or '无扩展名'}")


def _read_docx(path: Path) -> SourceDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("读取 Word 需要 python-docx，请安装项目的 lesson-plan 依赖") from exc

    document = Document(str(path))
    title = path.stem
    sections: list[SourceSection] = []
    current = SourceSection(title="课程导入", level=1)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if style_name == "title":
            title = text
            continue
        level = _heading_level(style_name)
        if level is not None:
            if current.paragraphs or current.title != "课程导入":
                sections.append(current)
            current = SourceSection(title=text, level=level)
        else:
            current.paragraphs.append(text)

    if current.paragraphs or not sections:
        sections.append(current)

    return SourceDocument(
        title=title,
        source_path=path,
        sections=sections,
        metadata={"format": "docx"},
    )


def _heading_level(style_name: str) -> int | None:
    for prefix in ("heading ", "标题 "):
        if style_name.startswith(prefix):
            try:
                return max(1, int(style_name[len(prefix):]))
            except ValueError:
                return 1
    return None


def _read_pdf(path: Path) -> SourceDocument:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError("读取 PDF 需要 pdfplumber，请安装项目的 lesson-plan 依赖") from exc

    sections = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            text = (page.extract_text() or "").strip()
            if text:
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                sections.append(
                    SourceSection(
                        title=f"第 {page_number} 页",
                        level=1,
                        paragraphs=lines,
                    )
                )

    if not sections:
        raise ValueError("PDF 未提取到文本；扫描版 PDF 需要先配置 OCR")

    return SourceDocument(
        title=path.stem,
        source_path=path,
        sections=sections,
        metadata={"format": "pdf", "page_count": len(sections)},
    )
